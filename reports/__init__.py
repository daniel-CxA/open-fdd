import math
import os
import time
from io import BytesIO

import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches

from faults import FaultConditionOne


class FaultCodeOneReport:
    """Class provides the definitions for Fault Code 1 Report."""

    def __init__(
        self,
        vfd_speed_percent_err_thres: float,
        vfd_speed_percent_max: float,
        duct_static_inches_err_thres: float,
        duct_static_col: str,
        supply_vfd_speed_col: str,
        duct_static_setpoint_col: str,
    ):
        self.vfd_speed_percent_err_thres = vfd_speed_percent_err_thres
        self.vfd_speed_percent_max = vfd_speed_percent_max
        self.duct_static_inches_err_thres = duct_static_inches_err_thres
        self.duct_static_col = duct_static_col
        self.supply_vfd_speed_col = supply_vfd_speed_col
        self.duct_static_setpoint_col = duct_static_setpoint_col

    def create_fan_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:
        if output_col is None:
            output_col = "fc1_flag"
            
        df[output_col] = df[output_col].astype(int)
        
        fig, (ax1, ax2, ax3) = plt.subplots(3, 1, figsize=(25, 8))
        plt.title('Fault Conditions 1 Plot')

        ax1.plot(df.index, df[self.duct_static_col], label="STATIC")
        ax1.legend(loc='best')
        ax1.set_ylabel("Inch WC")
        
        ax2.plot(df.index, df[self.supply_vfd_speed_col], color="g", label="FAN")
        ax2.legend(loc='best')
        ax2.set_ylabel('%')

        ax3.plot(df.index, df[output_col], label="Fault", color="k")
        ax3.set_xlabel('Date')
        ax3.set_ylabel('Fault Flags')
        ax3.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc1_flag"
        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)
        # print("DAYS ALL DATA: ", total_days)
        total_hours = delta.sum() / pd.Timedelta(hours=1)
        # print("TOTAL HOURS: ", total_hours)
        hours_fc1_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)
        # print("FALT FLAG TRUE TOTAL HOURS: ", hours_fc1_mode)
        percent_true = round(df[output_col].mean() * 100, 2)
        # print("PERCENT TIME WHEN FLAG IS TRUE: ", percent_true, "%")
        percent_false = round((100 - percent_true), 2)
        # print("PERCENT TIME WHEN FLAG 5 FALSE: ", percent_false, "%")
        flag_true_duct_static = round(
            df[self.duct_static_col].where(df[output_col] == 1).mean(), 2
        )
        return (
            total_days,
            total_hours,
            hours_fc1_mode,
            percent_true,
            percent_false,
            flag_true_duct_static,
        )

    def create_hist_plot(
        self, df: pd.DataFrame, output_col: str = None, duct_static_col: str = None
    ) -> plt:
        if output_col is None:
            output_col = "fc1_flag"
        if duct_static_col is None:
            duct_static_col = "duct_static"
        # calculate dataset statistics
        df["hour_of_the_day_fc1"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc3
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc1.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 1 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None,
        duct_static_col: str = None,
        flag_true_duct_static: bool = None,
    ) -> None:
        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition One Report", 0)

        p = document.add_paragraph(
            """Fault condition one of ASHRAE Guideline 36 is related to flagging poor performance of a AHU variable supply fan attempting to control to a duct pressure setpoint. Fault condition equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc1_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_fan_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)
        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc1_mode,
            percent_true,
            percent_false,
            flag_true_duct_static,
        ) = self.summarize_fault_times(df, output_col=output_col)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc1_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()
        # ADD HIST Plots
        document.add_heading("Time-of-day Histogram Plots", level=2)
        histogram_plot_image = BytesIO()
        histogram_plot = self.create_hist_plot(
            df, output_col=output_col, duct_static_col=duct_static_col
        )
        histogram_plot.savefig(histogram_plot_image, format="png")
        histogram_plot_image.seek(0)
        document.add_picture(
            histogram_plot_image,
            width=Inches(6),
        )

        if not math.isnan(flag_true_duct_static):
            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f'Average duct system pressure for when in fault condition (fan VFD speed > 95%): {flag_true_duct_static}"WC'
            )

        paragraph = document.add_paragraph()

        # ADD in Summary Statistics of fan operation
        document.add_heading("VFD Speed Statistics", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(str(df[self.supply_vfd_speed_col].describe()))

        # ADD in Summary Statistics of duct pressure
        document.add_heading("Duct Pressure Statistics", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(str(df[self.duct_static_col].describe()))

        # ADD in Summary Statistics of duct pressure
        document.add_heading("Duct Pressure Setpoints Statistics", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(str(df[self.duct_static_setpoint_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true > 5.0:
            paragraph.add_run(
                "The percent True metric that represents the amount of time for when the fault flag is True is high indicating the fan is running at high speeds and appearing to not generate good duct static pressure"
            )

        else:
            paragraph.add_run(
                "The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the fan appears to generate good duct static pressure"
            )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"


        if df[self.duct_static_setpoint_col].std() == 0:
            paragraph.add_run("No duct pressure setpoint reset detected (BAD)")

        else:
            paragraph.add_run("Duct pressure reset detected (Good)")

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document



class FaultCodeTwoReport:
    """Class provides the definitions for Fault Code 2 Report."""

    def __init__(
        self,
        mix_degf_err_thres: float,
        return_degf_err_thres: float,
        outdoor_degf_err_thres: float,
        mat_col: str,
        rat_col: str,
        oat_col: str,
    ):
        self.mix_degf_err_thres = mix_degf_err_thres
        self.return_degf_err_thres = return_degf_err_thres
        self.outdoor_degf_err_thres = outdoor_degf_err_thres
        self.mat_col = mat_col
        self.rat_col = rat_col
        self.oat_col = oat_col

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:
        if output_col is None:
            output_col = "fc2_flag"
            
        df[output_col] = df[output_col].astype(int)
        
        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(25, 8))
        plt.title('Fault Conditions 2 Plot')

        plot1a, = ax1.plot(df.index, df[self.mat_col], 
                           color='r', label="Mix Temp")  # red
        
        plot1b, = ax1.plot(df.index, df[self.rat_col], 
                           color='b', label="Return Temp") # blue
        
        plot1c, = ax1.plot(df.index, df[self.oat_col], 
                           color='g', label="Out Temp")  # green
        
        ax1.legend(loc='best')
        ax1.set_ylabel("°F")

        ax2.plot(df.index, df[output_col], label="Fault", color="k")
        ax2.set_xlabel('Date')
        ax2.set_ylabel('Fault Flags')
        ax2.legend(loc='best')

        plt.legend()
        plt.tight_layout()
        
        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc2_flag"
        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)
        # print("DAYS ALL DATA: ", total_days)
        total_hours = delta.sum() / pd.Timedelta(hours=1)
        # print("TOTAL HOURS: ", total_hours)
        hours_fc2_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)
        # print("FALT FLAG TRUE TOTAL HOURS: ", hours_fc1_mode)
        percent_true = round(df[output_col].mean() * 100, 2)
        # print("PERCENT TIME WHEN FLAG IS TRUE: ", percent_true, "%")
        percent_false = round((100 - percent_true), 2)
        # print("PERCENT TIME WHEN FLAG 5 FALSE: ", percent_false, "%")
        
        flag_true_mat = round(
            df[self.mat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_oat = round(
            df[self.oat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_rat = round(
            df[self.rat_col].where(df[output_col] == 1).mean(), 2
        )
        return (
            total_days,
            total_hours,
            hours_fc2_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_oat,
            flag_true_rat
        )

    def create_hist_plot(
        self, df: pd.DataFrame, 
        output_col: str = None, 
        mat_col: str = None
    ) -> plt:
        
        if output_col is None:
            output_col = "fc2_flag"
            
        if mat_col is None:
            mat_col = "mat"
            
        # calculate dataset statistics
        df["hour_of_the_day_fc2"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc3
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc2.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 2 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None,
        mat_col: str = None,
        flag_true_mat: bool = None,
    ) -> None:
        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Two Report", 0)

        p = document.add_paragraph(
            """Fault condition two and three of ASHRAE Guideline 36 is related to flagging mixing air temperatures of the AHU that are out of acceptable ranges. Fault condition 2 flags mixing air temperatures that are too low and fault condition 3 flags mixing temperatures that are too high when in comparision to return and outside air data. The mixing air temperatures in theory should always be in between the return and outside air temperatures ranges. Fault condition two equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc2_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)
        
        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc2_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_oat,
            flag_true_rat
        ) = self.summarize_fault_times(df, output_col=output_col)
        
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc2_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()
        # ADD HIST Plots
        document.add_heading("Time-of-day Histogram Plots", level=2)
        histogram_plot_image = BytesIO()
        histogram_plot = self.create_hist_plot(
            df, output_col=output_col, mat_col=mat_col
        )
        histogram_plot.savefig(histogram_plot_image, format="png")
        histogram_plot_image.seek(0)
        document.add_picture(
            histogram_plot_image,
            width=Inches(6),
        )

        if not math.isnan(flag_true_mat):
            paragraph = document.add_paragraph()
            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 2 is True the average mix air temp is {flag_true_mat}°F, outside air temp is {flag_true_oat}°F, and return air temp is {flag_true_rat}°F. This could possibly help with pin pointing AHU operating conditions for when this fault is True.')


        paragraph = document.add_paragraph()

        # ADD in Summary Statistics
        document.add_heading('Mix Temp Statistics', level=2)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df[self.mat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Return Temp Statistics', level=2)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df[self.rat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Outside Temp Statistics', level=2)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df[self.oat_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
            
            
        if percent_true < 5:

            paragraph.add_run(
                'The percent True of time in fault condition 2 or 3 is high indicating the AHU temperature sensors are out of calibration')

        else:
            paragraph.add_run(
                'The percent True of time is low inidicating the AHU temperature sensors are within calibration')


        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document
